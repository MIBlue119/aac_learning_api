import json
import io
import streamlit as st
from reportlab.platypus import SimpleDocTemplate
from reportlab.lib.pagesizes import letter
from reportlab.platypus import PageBreak
from aac_assets_generator.learning_asset_models import LearningAsset
from aac_assets_generator.learning_evaluation_models import EvaluationAssetTable
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def extract_main_title(prompt_content):
    pattern = r'([\u4e00-\u9fff]+系列)(?=的)'
    match = re.search(pattern, prompt_content)
    if match:
        return match.group(1)[2:]
    return "AAC系列"

async def get_user_study_sheet_data_async(session, api_key):
    url = "https://aaclearningbackend.azurewebsites.net/api/WebAAC/GetUserStudySheetData"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    async with session.get(url, headers=headers) as response:
        if response.status == 200:
            return await response.json()
        else:
            raise Exception(f"GetUserStudySheetData API 調用失敗，狀態碼 {response.status}")


async def get_board_prompt_word_data_async(session, api_key, board_id):
    url = "https://aaclearningbackend.azurewebsites.net/api/WebAAC/GetBoardPromptWordData"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"ID": board_id}
    async with session.get(url, headers=headers, data=json.dumps(data)) as response:
        if response.status == 200:
            return await response.json()
        else:
            raise Exception(f"GetBoardPromptWordData API 調用失敗，狀態碼 {response.status}")


def parse_user_data(user_data):
    def parse_json_field(field):
        if field:
            try:
                return ", ".join(json.loads(field))
            except json.JSONDecodeError:
                return field
        return "未提供"

    name = parse_json_field(user_data.get("name", "未提供"))
    gender = parse_json_field(user_data.get("gender", "未提供"))
    disability = parse_json_field(user_data.get("disability", "未提供"))
    communication_issues = parse_json_field(user_data.get("communication_Issues", "未提供"))
    communication_methods = parse_json_field(user_data.get("communication_Methods", "未提供"))
    strengths = parse_json_field(user_data.get("strengths", "未提供"))
    weaknesses = parse_json_field(user_data.get("weaknesses", "未提供"))
    teaching_time = parse_json_field(user_data.get("teaching_Time", "未提供"))

    case_info = f"""
    姓名: {name}
    性別: {gender}
    障礙類別: {disability}
    溝通問題: {communication_issues}
    溝通方式: {communication_methods}
    優勢能力: {strengths}
    弱勢能力: {weaknesses}
    預計教學時間: {teaching_time} 分鐘
    """

    return case_info


def combine_pdf_buffers(asset_elements, evaluate_elements):
    combined_elements = asset_elements + evaluate_elements
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    doc.build(combined_elements)
    
    return buffer

def export_assets_pdf(buffer, main_title, sub_title):
    st.download_button(
        label="下載 PDF",
        data=buffer,
        file_name=f"{main_title}-{sub_title}.pdf",
        mime="application/pdf",
        key="pdf_download"  
    )

def generate_combined_docx(learning_asset: LearningAsset, learning_evaluate: EvaluationAssetTable, main_title, sub_title):
    doc = Document()
    
    # Set font for the entire document
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    # Add title
    doc.add_heading(f"{main_title}-{sub_title}", level=0)    
    
    # Add title
    doc.add_heading('教案', level=0)
    # Add lesson plan table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Fill the table
    rows = table.rows
    rows[0].cells[0].text = '教案名稱'
    rows[0].cells[1].text = learning_asset.lesson_plan.title
    rows[1].cells[0].text = '教學目標'
    rows[1].cells[1].text = learning_asset.lesson_plan.objectives
    rows[2].cells[0].text = '教學內容'
    rows[2].cells[1].text = '\n'.join([f"{i+1}. {content}" for i, content in enumerate(learning_asset.lesson_plan.content)])
    rows[3].cells[0].text = '教學方法'
    rows[3].cells[1].text = '\n'.join([f"{i+1}. {method.title}: {method.explanation}" for i, method in enumerate(learning_asset.lesson_plan.teaching_methods)])
    rows[4].cells[0].text = '教學步驟'
    rows[4].cells[1].text = '\n'.join([f"{i+1}. {step.title}: {step.explanation}" for i, step in enumerate(learning_asset.lesson_plan.teaching_steps)])
    rows[5].cells[0].text = '評量方式'
    rows[5].cells[1].text = '\n'.join([f"{i+1}. {method.title}: {method.explanation}" for i, method in enumerate(learning_asset.lesson_plan.assessment_methods)])
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(15)
    doc.add_page_break()
    # Add worksheet
    doc.add_heading('學習單', level=1)
    doc.add_heading('一、練習題', level=2)
    for i, question in enumerate(learning_asset.worksheet.practice_questions, 1):
        doc.add_paragraph(f"{i}. {question.question}")
    doc.add_heading('二、活動指導', level=2)
    for i, guide in enumerate(learning_asset.worksheet.activity_guides, 1):
        doc.add_paragraph(f"{i}. {guide.description}")
    doc.add_heading('三、反思問題', level=2)
    for i, question in enumerate(learning_asset.worksheet.reflection_questions, 1):
        doc.add_paragraph(f"{i}. {question.question}")
    doc.add_heading('四、評量題', level=2)
    for i, question in enumerate(learning_asset.worksheet.assessment_questions, 1):
        doc.add_paragraph(f"{i}. {question.question}")
    doc.add_heading('五、自我評估表', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '評估項目'
    hdr_cells[1].text = '滿意(✓)'
    hdr_cells[2].text = '需改進(✗)'
    hdr_cells[3].text = '反思與改進方法'
    for item in learning_asset.worksheet.self_assessment_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.item
    doc.add_heading('六、合作學習活動', level=2)
    doc.add_paragraph(learning_asset.worksheet.collaborative_learning_activity)
    doc.add_page_break()
    # Add title
    doc.add_heading(learning_evaluate.evaluation_asset_title, level=1)
    # Add evaluation table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '評量項目'
    hdr_cells[1].text = '評量指標'
    hdr_cells[2].text = '優良（4分）'
    hdr_cells[3].text = '良好（3分）'
    hdr_cells[4].text = '尚可（2分）'
    hdr_cells[5].text = '待加強（1分）'
    for item in learning_evaluate.evaluation_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.evaluation_item_title
        row_cells[1].text = item.evaluation_metric
        row_cells[2].text = item.score_descriptions.excellent_with_score_4
        row_cells[3].text = item.score_descriptions.good_with_score_3
        row_cells[4].text = item.score_descriptions.fair_with_score_2
        row_cells[5].text = item.score_descriptions.needs_improvement_with_score_1
    # Set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(3)
    doc.add_paragraph()  # Add some space
    # Add scoring criteria
    doc.add_heading('評分標準', level=2)
    number_of_evaluation_items = len(learning_evaluate.evaluation_items)
    criteria = [
        f"優良: {3*(number_of_evaluation_items-1)}-{4*number_of_evaluation_items} 分, 表示學生能充分掌握技巧並理解其重要性。",
        f"良好: {2*(number_of_evaluation_items)}-{3*(number_of_evaluation_items-1)} 分, 表示學生能較好地完成步驟，但仍有待改進的部分。",
        f"尚可: {1*(number_of_evaluation_items)}-{2*(number_of_evaluation_items-1)} 分, 表示學生能完成部分步驟，但正確性和時間效率需加強。",
        f"待加強: {1*(number_of_evaluation_items-1)} 分, 表示學生需更多練習和輔助以掌握技巧。"
    ]
    for criterion in criteria:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(criterion)
    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file 

def export_asset_docx(docx_buffer,  main_title, sub_title):
    st.download_button(
           label="下載 Word 文件",
           data=docx_buffer.getvalue(),
           file_name=f"{main_title}-{sub_title}.docx",
           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
           key="docx_download"  # 添加唯一的 key
       )               


def render_streamlit_interface(learning_asset, learning_evaluate, asset_elements, evaluate_elements):
    # 初始化 session_state
    if 'pdf_buffer' not in st.session_state:
        st.session_state.pdf_buffer = None
    if 'docx_buffer' not in st.session_state:
        st.session_state.docx_buffer = None

    # 生成 PDF
    if st.session_state.pdf_buffer is None:
        st.session_state.pdf_buffer = combine_pdf_buffers(
            asset_elements,
            evaluate_elements
        )

    # 生成 DOCX
    if st.session_state.docx_buffer is None:
        st.session_state.docx_buffer = generate_combined_docx(learning_asset, learning_evaluate)


    st.subheader("下載 PDF 版本")
    export_assets_pdf(st.session_state.pdf_buffer)

    st.subheader("下載 Word 版本")
    export_asset_docx(st.session_state.docx_buffer)